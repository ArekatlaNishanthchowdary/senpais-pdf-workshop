"""PDF to Word via pdfplumber's positioned text, not a lossy PDF-import
round-trip. Good for text-heavy documents -- headings are guessed from line
height relative to the body-text height. Multi-column layouts, rotated text,
and graphic-heavy pages will still degrade; this reconstructs a reading order,
not a pixel-perfect clone. Kept out of the default install like OCR/tables --
see the `extras` group in pyproject.toml.
"""

from __future__ import annotations

import importlib
from collections import Counter
from pathlib import Path

from ..pages import output_path
from ..registry import register


@register(
    id="pdf_to_word",
    label="PDF to Word",
    category="Convert",
    summary="Reconstruct a PDF's text as real, editable Word paragraphs and headings.",
)
def pdf_to_word(sources: list[Path], out_dir: Path) -> list[Path]:
    try:
        import pdfplumber
        from docx import Document
    except ImportError as exc:
        raise ValueError(
            'PDF to Word needs the extras install: pip install "quire[extras]".'
        ) from exc

    # ponytail: pdfminer.six's standard-14 font metrics table (pdfminer.pdffont
    # .FONT_METRICS) can get corrupted by ocrmypdf's own pdfminer usage running
    # earlier in this same process -- every "Helvetica" font afterward silently
    # decodes to "(cid:N)" placeholders instead of real characters, no error
    # raised. Reloading the module resets that table; ceiling: if a future
    # pdfminer.six version keys this differently, this stops helping and needs
    # re-diagnosing rather than blindly kept around.
    import pdfminer.pdffont

    importlib.reload(pdfminer.pdffont)

    src = sources[0]
    with pdfplumber.open(str(src)) as pdf:
        pages = [page.extract_text_lines(return_chars=False) for page in pdf.pages]

    heights = [
        round(line["bottom"] - line["top"])
        for lines in pages
        for line in lines
        if line["text"].strip()
    ]
    if not heights:
        raise ValueError("No extractable text found in this document.")
    # ponytail: font size is the whole heuristic -- the most common line height
    # on the page is "body text"; noticeably taller lines are headings. No
    # attempt to read actual style/weight metadata.
    body_height = Counter(heights).most_common(1)[0][0]

    document = Document()
    for page_index, lines in enumerate(pages):
        if page_index > 0:
            document.add_page_break()
        for line in lines:
            text = line["text"].strip()
            if not text:
                continue
            height = line["bottom"] - line["top"]
            if height >= body_height * 1.8:
                document.add_heading(text, level=1)
            elif height >= body_height * 1.3:
                document.add_heading(text, level=2)
            else:
                document.add_paragraph(text)

    target = output_path(out_dir, src.stem, ".docx")
    document.save(target)
    return [target]
