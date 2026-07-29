"""PDF to Word via pdfplumber's positioned text, not a lossy PDF-import
round-trip. Good for text-heavy documents -- headings are guessed from line
height relative to the body-text height. Multi-column layouts, rotated text,
and graphic-heavy pages will still degrade; this reconstructs a reading order,
not a pixel-perfect clone. python-docx (the .docx writer) is kept out of the
default install, like OCR/tables -- see the `extras` group in pyproject.toml.
pdfplumber itself is a core dependency (also used by `redact`).
"""

from __future__ import annotations

from collections import Counter
from pathlib import Path

import pdfplumber

from ._pdfminer_fix import reset_pdfminer_font_cache
from ..pages import output_path
from ..registry import register


@register(
    id="pdf_to_word",
    label="PDF to Word",
    category="Convert",
    summary="Reconstruct a PDF's text as real, editable Word paragraphs and headings.",
)
def pdf_to_word(sources: list[Path], out_dir: Path) -> list[Path]:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError(
            'PDF to Word needs the extras install: pip install "senpais-pdf-workshop[extras]".'
        ) from exc

    reset_pdfminer_font_cache()

    src = sources[0]
    with pdfplumber.open(str(src)) as pdf:
        pages = [page.extract_text_lines(return_chars=False) for page in pdf.pages]

    heights = [
        round(line["bottom"] - line["top"])
        for lines in pages
        for line in lines
        if line["text"].strip()
    ]
    if not heights:
        raise ValueError("No extractable text found in this document.")
    # ponytail: font size is the whole heuristic -- the most common line height
    # on the page is "body text"; noticeably taller lines are headings. No
    # attempt to read actual style/weight metadata.
    body_height = Counter(heights).most_common(1)[0][0]

    document = Document()
    for page_index, lines in enumerate(pages):
        if page_index > 0:
            document.add_page_break()
        for line in lines:
            text = line["text"].strip()
            if not text:
                continue
            height = line["bottom"] - line["top"]
            if height >= body_height * 1.8:
                document.add_heading(text, level=1)
            elif height >= body_height * 1.3:
                document.add_heading(text, level=2)
            else:
                document.add_paragraph(text)

    target = output_path(out_dir, src.stem, ".docx")
    document.save(target)
    return [target]
